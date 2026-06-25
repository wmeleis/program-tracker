"""Generate a clean, shareable Word (.docx) document from a parsed reference
curriculum.

The custom_references table stores each reference as `sections_json` — a list of:

    {"heading": "<section name>",
     "courses": [
        {"is_header": true,  "text": "<sub-header or instruction>", "level": "block|area|inst"},
        {"is_header": false, "code": "BINF 6200", "title": "...", "hours": "4"},
        ...
     ]}

`level` is optional. When present it controls styling precisely; when absent
(e.g. parser-produced references) it's inferred heuristically. Output layout:

  • document title (Heading 0)
  • "block" headers  -> Heading 1   (e.g. concentration names, "Core Requirements")
  • "area" headers   -> bold line    (e.g. "Business Management", "Technical")
  • "inst" headers   -> italic line   (e.g. "Complete two of the following: (6 SH)")
  • a "Total credit hours: …" header -> bold summary line at the bottom
  • course rows       -> 3-column table (Course / Title / Hours)

This is the same content that drives the dashboard Compare/Reference tabs, so
clicking Download yields the shareable document — one source of truth.
"""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

_BOILERPLATE_HEADINGS = {
    'catalog presentation of this program',
    'course list',
}

# A header row that reads as an instruction (rendered italic, above its table)
# rather than a section/area label.
_INST_RE = re.compile(
    r'(following:|prerequisite|in place of'
    r'|^\s*complete\b|^\s*select\b|^\s*registration\b|^\s*then complete'
    r'|^\s*note\b|^\s*a concentration is required|^\s*thesis topic approval'
    r'|\bSH\))',
    re.I,
)
# A header row that should be a top-level Heading (concentration block, etc.).
_BLOCK_RE = re.compile(
    r'(concentration\s*$|without concentration|^\s*core requirements\s*$'
    r'|^\s*pathway differences)',
    re.I,
)


def _norm(h):
    return re.sub(r'\s+', ' ', (h or '').replace('\xa0', ' ')).strip()


def _classify(text, level):
    """Return 'block' | 'area' | 'inst' | 'total' for a header row."""
    if level in ('block', 'area', 'inst'):
        return level
    t = _norm(text)
    if re.match(r'(?i)^total credit hours', t):
        return 'total'
    if _INST_RE.search(t):
        return 'inst'
    if _BLOCK_RE.search(t):
        return 'block'
    return 'area'


def _course_sig(courses):
    return tuple(
        ((c.get('code') or '').strip().lower(), (c.get('title') or '').strip().lower())
        for c in (courses or []) if not c.get('is_header')
    )


def _sp(p, before, after):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return p


def build_reference_docx(name, sections, notes='', title=''):
    """Return .docx bytes for a parsed reference curriculum."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    heading_text = (title or name or 'Reference Curriculum').strip()
    doc.add_heading(heading_text, level=0)
    if notes:
        p = doc.add_paragraph()
        r = p.add_run(_norm(notes))
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _sp(p, 0, 8)

    tbl = [None]   # current open course table (mutable holder)

    def new_table():
        t = doc.add_table(rows=0, cols=3)
        try:
            t.style = 'Light Grid Accent 1'
        except Exception:
            t.style = 'Table Grid'
        return t

    def add_course(code, title_, hours):
        if tbl[0] is None:
            tbl[0] = new_table()
        cells = tbl[0].add_row().cells
        cells[0].paragraphs[0].add_run(code or '')
        cells[1].paragraphs[0].add_run(title_ or '')
        cells[2].paragraphs[0].add_run(hours or '')
        for c, w in zip(cells, (1.5, 4.7, 0.5)):
            c.width = Inches(w)
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    def add_header(text, kind):
        tbl[0] = None  # close any open table
        text = _norm(text)
        if not text:
            return
        if kind == 'block':
            _sp(doc.add_heading(text, level=1), 18, 4)
        elif kind == 'inst':
            p = doc.add_paragraph()
            p.add_run(text).italic = True
            _sp(p, 2, 2)
        elif kind == 'total':
            p = doc.add_paragraph()
            p.add_run(text).bold = True
            _sp(p, 18, 0)
        else:  # area
            p = doc.add_paragraph()
            p.add_run(text).bold = True
            _sp(p, 10, 2)

    seen = set()
    prev_heading = None
    for section in (sections or []):
        sec_heading = _norm(section.get('heading'))
        courses = section.get('courses') or []
        hl = sec_heading.lower()
        sig = _course_sig(courses)
        if (hl, sig) in seen and sig:
            continue
        seen.add((hl, sig))
        if sec_heading and hl not in _BOILERPLATE_HEADINGS and hl != prev_heading:
            add_header(sec_heading, _classify(sec_heading, None))
        prev_heading = hl
        for c in courses:
            if c.get('is_header'):
                add_header(c.get('text'), _classify(c.get('text'), c.get('level')))
            else:
                add_course((c.get('code') or '').strip(),
                           (c.get('title') or '').strip(),
                           (c.get('hours') or '').strip())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
